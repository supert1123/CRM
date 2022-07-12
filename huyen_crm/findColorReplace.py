from docxPdfImage import *
from docx.enum.text import WD_COLOR_INDEX
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""
Testing iter_block_items()
"""
def iter_block_items(parent):
    """
    Tạo tham chiếu đến từng đoạn và bảng con trong file doc, theo thứ tự tài liệu. 
    Mỗi giá trị trả về là một thể hiện của Bảng hoặc Đoạn văn. 
    'parent' thường là một tham chiếu đến một chính Đối tượng tài liệu, 
    hoạt động cho đối tượng _Cell | đoạn văn | bảng
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def color_string(key,countKey,p1,p):
    ##    tô vàng key
##    input: key, số thứ tự key, đoạn văn chứa key, đoạn mới chứa key được tô vàng
##    output: đoạn văn đã được tô vàng, số thứ tự key
    str1 =""
    sub_end =""
    #print(p1,"\n")
    for i in range(len(key)):
        substrings = p1.split(key[i]) # split đoạn
        if(i!=len(key)-1):
            p1 = str1.join(substrings)
    #print(substrings)  
        for substring in substrings[:-1]:
            #print('subs',substring)
            countKey += 1
            p.add_run(substring,style = 'CommentsStyle') # Ép kiểu chữ theo font'CommentStyle'
            font = p.add_run(key[i],style = 'CommentsStyle').font.highlight_color = WD_COLOR_INDEX.YELLOW # tô vàng key
            count = str(countKey)
            font = p.add_run(count,style = 'CommentsStyle').font.highlight_color = WD_COLOR_INDEX.RED # tô đỏ số thứ tự của key
        sub_end = substrings[-1]
    p.add_run(substrings[-1], style = 'CommentsStyle')
    return countKey
def Size(filename):# tìm và chọn size của văn bản
    size = []
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        for i in p.runs:
            if i.font.size != None:
                size.append(i.font.size/12700)
    return size
def iter_unique_cells(row): #(Hợp nhất cells theo dòng, bỏ qua các lần lặp lại)
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
def findColor(filename,key,newName):
    ##    tìm và tô vàng key
##    input: file cần xử ký, key cần tìm và tô vàng
##    output: file đã tô vàng và đánh thứ tự cho key
    countKey = 0 # khởi tạo số thứ tự key
    doc = Document(filename)
    
    #Tạo font theo 'CommentsStyle'
    par = doc.paragraphs[0]
    font_styles = doc.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    if '{par.style.font.name}'!= None:
        font_object.name = f"'{par.style.font.name}'"
    sizes = Size(filename)
    run = 0
    for p in doc.paragraphs:
        for i in p.runs:
            if i.font.size != None:
                font_object.size = Pt(sizes[run])
                run += 1
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            #print(type(block)
            p1 = block.text
            match = re.findall(key,p1,re.IGNORECASE)
            #print(match)
            #for igkey in matc
            if len(match)>0: #so khớp không phân biệt hoa thường
                block.text = ""
                countKey=color_string(match,countKey,p1,block)
        else:
            for row in block.rows:
                for p in iter_unique_cells(row):
                    p1 = p.text
                    match = re.findall(key,p1,re.IGNORECASE)
                    if len(match)>0: #so khớp không phân biệt hoa thường
                        p.text = " "
                        p = p.add_paragraph()
                        countKey = color_string(match,countKey,p1,p)
    doc.save(newName)
    return countKey

# def replace_string(key,value,numberList,countKey,p):
"""     
{
    'chiết khấu':{
        'abc': [1, 2 ,3],
        'xyz': [4,5,6]
    }
}
"""
def replace_string(dict_key, countKey,p):
    """    
    split đoạn văn và key thành list
    kiểm tra xem key có xuất hiện trong đoạn không
    đếm số lần xuất hiện của key, nếu thứ tự nằm trong numberList thì thay đổi key
    thay đổi: xóa các từ khác, giữ lại từ đầu tiên của key, thay bằng value
    chuyển đoạn văn từ list về đoạn

    input: key, từ để đổi, danh sách vị trí đổi, số thứ tự key, đoạn văn chứa key
    output: đoạn văn đã được đổi từ ở vị trí chỉ định, số thứ tự key
    """
    numberList = [ ]
    for k in dict_key:
        key = k
        for v in dict_key[k]:
            for numList in dict_key[k][v]:
                numberList.append(numList)
    line_split = p.text.split() # split đoạn
    key_split = key.split() # split key
    len_key = len(key_split)
    for i in range(len(line_split)):
        if re.findall(key_split[0],line_split[i],re.IGNORECASE):# nếu từ đầu trong key xuất hiện
            count = 0 # đếm từ trong key
            while count < len_key:
                if re.findall(key_split[count],line_split[i+count],re.IGNORECASE): ##so khớp không phân biệt hoa thường
                    count+=1 # đếm xem các từ trong key xuất hiện đủ chưa
                else:
                    break
            if count == len_key: # nếu đủ
                countKey += 1
                #punctuation =""
                #if re.match(r'\S', p.text): #so khớp với ký tự không phải chữ
                    #punctuation = line_split[i+count-1][-1] # dấu câu
                if countKey in numberList:  # bắt đầu thay đổi ở các vị trí cần thiết
                    count_1 = 1
                    while count_1 < len_key:
                        line_split[i+count_1] = "" #thêm u ở phía trước để xử lý ký tự tiếng việt nhá
                        count_1+=1
                    value = ''
                    for k in dict_key:
                        for v in dict_key[k]:
                            if countKey in dict_key[k][v]:
                                value = v
                    line_split[i] = value #+punctuation
                    # run = p.add_run()
                    # font = run.font                    
                    for idx in range(len(p.runs)): #khởi tạo idx
                        if idx<len(line_split)-1: # so sánh giá trị của idx
                            p.runs[idx].text = line_split[idx] + ' '  #tạo đoạn thay đổi thêm vào
                        elif idx==len(line_split)-1:      #so sánh thay đổi
                            p.runs[idx].text = u" ".join(line_split[idx:])   #Thay đổi
                            p.runs[idx].text = u" ".join(p.runs[idx].text.split()) 
                        else:  # Xét định dạng của đoạn thay đổi
                            p.runs[idx].text = ''     
                        flag = check_font(p.runs[idx])
                        if flag['bold']:    
                            p.runs[idx].font.bold = True
                        if flag['italic']:
                            p.runs[idx].font.italic = True
                        if flag['underline']:
                            p.runs[idx].font.underline = True
                    # run.text = u" ".join(line_split)
                    # a = run.text
                    # run.text = u" ".join(a.split()) #loai bỏ khoảng trắng trùng lặp   
                    # p.text = u" ".join(p.text.split())            
    return countKey

def check_font(para):
    flag = {
        'bold': 0,
        'italic':0,
        'underline':0,
    }
    if para.bold:
        flag['bold'] = 1
    if para.italic: 
        flag['italic'] = 1
    if para.underline: 
        flag['underline'] = 1
    return flag


def replace(filename,dict_key,output_file):
    """
    hàm duyệt từng đoạn trong file
    tìm và thay thế từ ở vị trí chỉ định
    input: tên file, từ muốn đổi, từ để đổi, danh sách vị trí đổi
    output: file word đã được thay từ ở những vị trí chỉ định
    """
    for d in dict_key:
        key = d
    countKey = 0 # khởi tạo số thứ tự key
    doc = Document(filename)
    par = doc.paragraphs[0]
    #Tạo Style của văn bản và  đưa dòng có key
    styles = doc.styles
    style = doc.styles['Normal']
    font = style.font
    if f"'{par.style.font.name}'" != None:
        font.name = f"'{par.style.font.name}'"
    sizes = Size(filename)
    run = 0
    for p in doc.paragraphs:
        for i in p.runs:
            if i.font.size != None:
                font.size = Pt(sizes[run])
                run += 1
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if re.findall(key,block.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                countKey = replace_string(dict_key,countKey,block)
        else:
            for row in block.rows:
                for cell in iter_unique_cells(row):
                    for p in cell.paragraphs:
                            if re.findall(key,p.text,re.IGNORECASE): #so khớp không phân biệt hoa thường
                                    #print(re.search(key,p.text,re.IGNORECASE))
                                    countKey = replace_string(dict_key,countKey,p)                                
                                    #Đưa style vào từ chuyển đổi   
                                    styles = doc.styles
                                    style = doc.styles['Normal']
                                    font = style.font
                                    font.name = f"'{par.style.font.name}'" 
                                    font.name = f"'{par.style.font.name}'" 
                                    font.name = f"'{par.style.font.name}'" 
    doc.save(output_file)

'''input_file = 'output/phong8.docx'
input_file = os.path.abspath(input_file)
#file = os.getcwd() + "/" + input_file
key = u'công việc'
value = u'công việc tuần'
numberList=[1,2]
output_file = 'output/output.docx'
replace(input_file,key,value,numberList,output_file)'''
